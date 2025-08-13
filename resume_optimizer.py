#!/usr/bin/env python3
"""
Resume Optimizer - Simple LLM workflow for ATS-optimized resume generation
"""

import os

from docx import Document
from dotenv import load_dotenv
from openai import OpenAI

# Load environment variables from parent directory
load_dotenv('../.env')
# load_dotenv()

class ResumeOptimizer:
    def __init__(self):
        print("🔄 Initializing OpenAI client...")
        api_key = os.getenv('OPEN_AI_KEY')
        if not api_key:
            raise ValueError("OPEN_AI_KEY not found in environment variables")
        self.client = OpenAI(api_key=api_key)
        print("✅ OpenAI client initialized successfully!")
    
    def read_files(self, jd_path, resume_path):
        """Read job description and resume files"""
        with open(jd_path, 'r') as f:
            jd_text = f.read()
        
        doc = Document(resume_path)
        resume_sections = [para.text for para in doc.paragraphs if para.text.strip()]
        resume_text = '\n'.join(resume_sections)
        
        return jd_text, resume_text, resume_sections
    
    def optimize_resume_complete(self, jd_text, resume_text):
        """Complete resume optimization in one LLM call"""
        prompt = f"""You are a senior resume writer with over 10 years of experience specializing in optimizing resumes for Applicant Tracking Systems (ATS). Your skills include keyword optimization, skills matching, and industry-specific formatting.

RULES:

Enhancement Only: Enhance existing experience without fabricating details.
Keyword Use: Use exact keywords from the job description where applicable.
Professional Tone: Maintain a professional tone and formatting throughout.
Quantifiable Achievements: Focus on quantifiable achievements whenever possible.
ATS-Friendly Format: Ensure the resume is ATS-friendly, avoiding tables, graphics, or complex layouts.

Task:
Analyze the job requirements and optimize the provided resume.

Provide:

KEY REQUIREMENTS: Extract and list the top 5 must-have requirements from the job posting.
OPTIMIZED RESUME: Deliver an enhanced version of the resume with strategic keyword placement and improved formatting.
IMPROVEMENTS MADE: Detail specific changes made to the resume and explain the rationale behind each modification.
FURTHER IMPROVEMENTS: Provide places that can be further improved upon with more evidence.

Inputs:

Job Description: {jd_text}
Current Resume: {resume_text}

Response:"""
        
        response = self.client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a professional resume writer and ATS expert."},
                {"role": "user", "content": prompt}
            ],
            max_completion_tokens=1500,
            temperature=0.5 #temperature is not supported...
        )
        return response.choices[0].message.content.strip()
    
    def optimize_from_text(self, jd_text, resume_text):
        """Optimize resume from text input (for web interface)"""
        llm_response = self.optimize_resume_complete(jd_text, resume_text)
        return self.parse_llm_response(llm_response)
    
    def parse_llm_response(self, response):
        """Parse the LLM response into components"""
        sections = {
            'requirements': 'Not provided', 
            'optimized_resume': '',
            'improvements': 'Not provided',
            'further_improvements': 'Not provided'
        }
        
        # More flexible parsing - look for various patterns
        import re
        
        # No ATS score extraction needed
        
        # Extract sections using more flexible splitting
        lines = response.split('\n')
        current_section = None
        content = []
        
        for line in lines:
            line = line.strip()
            if 'KEY REQUIREMENTS' in line.upper() or 'REQUIREMENTS' in line.upper():
                if content and current_section:
                    sections[current_section] = '\n'.join(content)
                current_section = 'requirements'
                content = []
            elif 'OPTIMIZED RESUME' in line.upper() or 'ENHANCED RESUME' in line.upper():
                if content and current_section:
                    sections[current_section] = '\n'.join(content)
                current_section = 'optimized_resume'
                content = []
            elif 'IMPROVEMENTS MADE' in line.upper() or 'CHANGES' in line.upper():
                if content and current_section:
                    sections[current_section] = '\n'.join(content)
                current_section = 'improvements'
                content = []
            elif 'FURTHER IMPROVEMENTS' in line.upper():
                if content and current_section:
                    sections[current_section] = '\n'.join(content)
                current_section = 'further_improvements'
                content = []
            elif line and current_section:
                content.append(line)
        
        # Don't forget the last section
        if content and current_section:
            sections[current_section] = '\n'.join(content)
        
        # If no optimized resume found, use the entire response
        if not sections['optimized_resume'] and len(response) > 100:
            sections['optimized_resume'] = response
        
        return sections
    
    def write_resume(self, optimized_text, output_path):
        """Write optimized resume to DOCX file"""
        doc = Document()
        
        # Split into paragraphs and add to document
        paragraphs = optimized_text.split('\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        doc.save(output_path)
    
    def run(self, jd_file='jd.txt', resume_file='resume.docx', output_file='new_resume.docx'):
        """Main workflow - single LLM call for everything"""
        try:
            print("📄 Reading files...")
            jd_text, resume_text, resume_sections = self.read_files(jd_file, resume_file)
            print(f"✅ Loaded JD and resume ({len(resume_sections)} sections)")
            
            print("🤖 Running AI optimization (this may take 30-60 seconds)...")
            llm_response = self.optimize_resume_complete(jd_text, resume_text)
            
            print("📊 Parsing results...")
            print(f"Debug - LLM Response: {llm_response[:500]}...")  # Show first 500 chars
            results = self.parse_llm_response(llm_response)
            
            if results['optimized_resume']:
                print("💾 Saving optimized resume...")
                self.write_resume(results['optimized_resume'], output_file)
                print(f"✅ New resume saved as: {output_file}")
            else:
                print("⚠️  No optimized resume generated, using original")
                self.write_resume(resume_text, output_file)
            
            # Display results
            print("\n" + "="*50)
            print("🎯 OPTIMIZATION RESULTS")
            print("="*50)

            print(f"🔑 Key Requirements:\n{results['requirements']}")
            print(f"🔧 Improvements Made:\n{results['improvements']}")
            print(f"📈 Further Improvements:\n{results['further_improvements']}")
            print("="*50)
            
            return results
            
        except FileNotFoundError as e:
            print(f"❌ Error: File not found - {e}")
            print("Make sure 'jd.txt' and 'resume.docx' exist in the current directory")
        except Exception as e:
            print(f"❌ Error: {e}")

if __name__ == "__main__":
    optimizer = ResumeOptimizer()
    optimizer.run()